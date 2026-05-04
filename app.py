import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

st.set_page_config(page_title="Sabaq AI Free", page_icon="📘", layout="wide")

st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg,#EEF2FF,#ECFEFF,#FFFFFF);
}
.hero {
    background: linear-gradient(135deg,#2563EB,#7C3AED,#EC4899);
    padding:35px;
    border-radius:30px;
    color:white;
    text-align:center;
    margin-bottom:25px;
    box-shadow:0 20px 50px rgba(37,99,235,.35);
}
.hero h1 {
    font-size:48px;
    font-weight:900;
}
.card {
    background:white;
    padding:25px;
    border-radius:25px;
    box-shadow:0 15px 35px rgba(0,0,0,.1);
}
.stButton > button {
    height:60px;
    border-radius:18px;
    font-size:18px;
    font-weight:800;
    color:white;
    background:linear-gradient(135deg,#2563EB,#7C3AED,#EC4899);
    border:none;
}
.stDownloadButton > button {
    border-radius:15px;
    font-weight:800;
    background:linear-gradient(135deg,#10B981,#06B6D4);
    color:white;
    border:none;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="hero">
<h1>📘 Sabaq AI FREE</h1>
<p>Тегін ҚМЖ генератор: ҚМЖ • Word кесте • Жұмыс парағы • Ойын идеялары</p>
</div>
""", unsafe_allow_html=True)

with st.sidebar:
    st.header("⚙️ Сабақ мәліметтері")
    school = st.text_input("ББҰ / мектеп атауы", "№1 Хромтау орта мектебі")
    teacher = st.text_input("Педагогтің Т.А.Ә.", "Сисекенова А.М.")
    subject = st.text_input("Пән", "Биология")
    grade = st.text_input("Сынып", "7")
    section = st.text_input("Бөлім", "")
    date = st.text_input("Күні", "")
    topic = st.text_area("Сабақ тақырыбы", "")
    objective = st.text_area("Оқу бағдарламасына сәйкес оқыту мақсаттары", "")
    lesson_type = st.selectbox("Сабақ түрі", ["Жаңа сабақ", "Ашық сабақ", "Зертханалық жұмыс", "Практикалық сабақ", "Қайталау сабағы"])
    ebq = st.radio("ЕБҚ тапсырмасы қосылсын ба?", ["Иә", "Жоқ"])

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tbl_pr.append(borders)

def add_p(doc, text, bold=False, size=11, align="left"):
    p = doc.add_paragraph()
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def qmj_data():
    lesson_goal = f"{topic} тақырыбы бойынша оқушылар негізгі ұғымдарды түсіндіреді, оқу мақсатына сай тапсырмаларды орындайды және алған білімін өмірлік жағдаятта қолданады."
    ebq_goal = f"{topic} бойынша негізгі тірек сөздерді пайдаланып, жеңілдетілген тапсырманы орындайды." if ebq == "Иә" else "Қарастырылмаған"

    flow = [
        ["Сабақтың басы\n5 минут", "Оқушылармен сәлемдеседі, түгендейді. Психологиялық ахуал қалыптастырады. Сабақ тақырыбы мен оқу мақсатын таныстырады.", "Сабаққа назар аударады, мақсатпен танысады.", "ҚБ: ауызша мадақтау.", "Презентация"],
        ["Өткен білімді еске түсіру\n5 минут", f"Өткен тақырыппен байланысты 3 сұрақ қояды. {topic} тақырыбына бағыттайды.", "Сұрақтарға жауап береді, өз ойын айтады.", "Дескриптор: өткен білімді еске түсіреді – 1 балл; жауапты дәлелдейді – 1 балл.", "Wordwall / карточка"],
        ["Жаңа тақырыпты ашу\n5 минут", f"Проблемалық сұрақ қояды: «{topic} не үшін маңызды?» Оқушылардың болжамын тыңдайды.", "Болжам жасайды, тақырыпты анықтайды.", "ҚБ: мұғалімнің кері байланысы.", "Сурет, бейнематериал"],
        ["Мұғалім түсіндірмесі\n7 минут", f"{topic} бойынша негізгі ұғымдарды түсіндіреді. Оқу мақсаты: {objective}. Терминдермен жұмыс жүргізеді.", "Негізгі ақпаратты дәптерге жазады.", "Дескриптор: негізгі ұғымды атайды – 1 балл; түсіндіреді – 1 балл.", "Презентация, оқулық"],
        ["Топтық жұмыс\n8 минут", f"1-топ: сызба құрастырады. 2-топ: түсіндіреді. 3-топ: өмірлік мысалмен байланыстырады.", "Топта талқылайды, постер/сызба жасайды, қорғайды.", "Дескриптор: мазмұнын ашады – 1 балл; дәлел келтіреді – 1 балл; қорытынды жасайды – 1 балл.", "Постер, маркер"],
        ["Жеке жұмыс\n5 минут", f"{topic} бойынша 3 сұрақтан тұратын жеке тапсырма орындайды.", "Жауап жазады, өз білімін қолданады.", "Дескриптор: дұрыс жауап береді – 1 балл; түсіндіреді – 1 балл.", "Жұмыс парағы"],
        ["PISA тапсырмасы\n5 минут", f"Жағдаят: оқушы күнделікті өмірде {topic} құбылысымен кездесті. Неліктен бұлай болды? Қандай қорытынды жасауға болады?", "Жағдаятты оқиды, жауап таңдайды, дәлелдейді.", "Дескриптор: жағдаятты түсінеді – 1 балл; дұрыс шешім ұсынады – 1 балл; дәлелдейді – 1 балл.", "PISA карточкасы"],
        ["ЕБҚ тапсырмасы\n3 минут", f"Тірек сөздер арқылы {topic} бойынша сәйкестендіру немесе ретімен орналастыру тапсырмасын орындайды.", "Мұғалім көмегімен жеңілдетілген тапсырма орындайды.", "Дескриптор: тірек сөзді таниды – 1 балл; дұрыс орналастырады – 1 балл.", "Карточка"],
        ["Сабақ соңы\n2 минут", "Рефлексия жүргізеді: «Бүгін білдім...», «Маған қиын болды...», «Мен қолдана аламын...».", "Өз ойын жазады немесе ауызша айтады.", "10 баллдық жүйе бойынша қорытынды бағалау.", "Стикер"]
    ]

    return {
        "school": school,
        "section": section,
        "teacher": teacher,
        "date": date,
        "grade": grade,
        "topic": topic,
        "objective": objective,
        "lesson_goal": lesson_goal,
        "ebq_goal": ebq_goal,
        "values": "Жауапкершілік, ынтымақтастық, академиялық адалдық, еңбекқорлық",
        "flow": flow
    }

def create_qmj_docx(data):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.3)
    sec.right_margin = Cm(1.3)

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)

    add_p(doc, "ҚЫСҚА МЕРЗІМДІ ЖОСПАР", True, 14, "center")

    info = [
        ("ББҰ", data["school"]),
        ("Бөлім", data["section"]),
        ("Педагогтің Т.А.Ә.", data["teacher"]),
        ("Күні", data["date"]),
        ("Сынып", data["grade"]),
        ("Сабақтың тақырыбы", data["topic"]),
        ("Оқу бағдарламасына сәйкес оқыту мақсаттары", data["objective"]),
        ("Сабақтың мақсаты", data["lesson_goal"]),
        ("ЕБҚЕ мақсаты", data["ebq_goal"]),
        ("Құндылықтар", data["values"]),
    ]

    t = doc.add_table(rows=len(info), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)

    for i, (a, b) in enumerate(info):
        set_cell_text(t.cell(i, 0), a, True, 9)
        set_cell_text(t.cell(i, 1), b, False, 9)
        set_cell_shading(t.cell(i, 0), "EAF2F8")

    doc.add_paragraph("")
    add_p(doc, "Сабақтың барысы", True, 13, "center")

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    headers = ["Сабақ кезеңі", "Педагогтің әрекеті", "Оқушының әрекеті", "Бағалау", "Ресурстар"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, True, 8)
        set_cell_shading(table.cell(0, i), "D9EAF7")

    for row_data in data["flow"]:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(row[i], value, i == 0, 8)

    doc.add_paragraph("")
    add_p(doc, "Үй тапсырмасы:", True, 11)
    add_p(doc, f"{topic} тақырыбын оқу, негізгі терминдерді қайталау, 3 сұрақ құрастыру.", False, 11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def worksheet_text():
    return f"""
ОҚУШЫҒА АРНАЛҒАН ЖҰМЫС ПАРАҒЫ

Пән: {subject}
Сынып: {grade}
Тақырып: {topic}
Оқу мақсаты: {objective}
Оқушының аты-жөні: ______________________
Күні: ______________________

1-тапсырма. Терминдерді сәйкестендір.
Терминдерді анықтамаларымен сәйкестендіріңіз.

2-тапсырма. Бос орынды толықтыр.
{topic} тақырыбы бойынша негізгі ұғымдарды қолданып, сөйлемдерді толықтырыңыз.

3-тапсырма. Сызба құрастыр.
Тақырып бойынша негізгі үдерісті немесе ұғымдар байланысын сызба түрінде көрсетіңіз.

4-тапсырма. PISA тапсырмасы.
Жағдаят: Күнделікті өмірде {topic} құбылысына байланысты жағдай орын алды.
Сұрақ: Бұл құбылыстың себебі қандай?
A) Кездейсоқ жағдай
B) Ғылыми заңдылыққа байланысты
C) Тек сыртқы әсер
Жауабыңызды дәлелдеңіз: ______________________

5-тапсырма. Жоғары деңгейлі сұрақ.
{topic} тақырыбының өмірдегі маңызын түсіндіріңіз.

ЕБҚ тапсырмасы:
Тірек сөздерді пайдаланып, тақырыптың негізгі идеясын жазыңыз.

Өзін-өзі бағалау:
Мен тақырыпты түсіндім: Иә / Жартылай / Қиын болды
Мен тапсырманы орындадым: Иә / Жартылай / Қиын болды

Рефлексия:
Бүгін мен білдім: ______________________
Маған қиын болды: ______________________
Мен үшін қызықты болды: ______________________
"""

def games_text():
    return f"""
САБАҚҚА АРНАЛҒАН ОЙЫН ИДЕЯЛАРЫ

Тақырып: {topic}

1. Wordwall
Ойын түрі: Сәйкестендіру
Тапсырма: {topic} бойынша термин мен анықтаманы сәйкестендіру.

2. Kahoot
Ойын түрі: Викторина
10 сұрақ құрастыруға болады:
- {topic} дегеніміз не?
- Бұл үдерістің маңызы қандай?
- Қай жауап дұрыс?
- Қай тұжырым жалған?

3. Genially
Формат: Интерактивті сурет
Оқушылар суреттегі белгілерді басып, түсініктеме ашады.

4. LearningApps
Ойын түрі: Реттілікке қою
Тапсырма: тақырыптағы кезеңдерді дұрыс ретімен орналастыру.

5. Офлайн ойын
Атауы: «Ғылыми детектив»
Оқушыларға жағдаят беріледі, олар себебін анықтап, дәлелдейді.
"""

def create_text_docx(text, title):
    doc = Document()
    add_p(doc, title, True, 14, "center")
    for line in text.split("\n"):
        add_p(doc, line, False, 11)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def check():
    if not topic or not objective:
        st.warning("Алдымен сабақ тақырыбы мен оқу мақсатын толтырыңыз.")
        return False
    return True

col1, col2, col3 = st.columns(3)

if col1.button("📘 ҚМЖ жасау + Word кесте", use_container_width=True):
    if check():
        data = qmj_data()
        st.success("✅ ҚМЖ дайын!")
        st.dataframe(data["flow"], use_container_width=True)
        st.download_button(
            "📥 ҚМЖ Word жүктеу",
            create_qmj_docx(data),
            file_name=f"QMJ_{topic[:20]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if col2.button("📄 Жұмыс парағын жасау", use_container_width=True):
    if check():
        text = worksheet_text()
        st.success("✅ Жұмыс парағы дайын!")
        st.markdown(f"<div class='card'>{text}</div>", unsafe_allow_html=True)
        st.download_button(
            "📥 Жұмыс парағын Word жүктеу",
            create_text_docx(text, "ЖҰМЫС ПАРАҒЫ"),
            file_name=f"Worksheet_{topic[:20]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if col3.button("🎮 Ойындар табу", use_container_width=True):
    if check():
        text = games_text()
        st.success("✅ Ойын идеялары дайын!")
        st.markdown(f"<div class='card'>{text}</div>", unsafe_allow_html=True)
        st.download_button(
            "📥 Ойындарды Word жүктеу",
            create_text_docx(text, "ОЙЫН ИДЕЯЛАРЫ"),
            file_name=f"Games_{topic[:20]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
